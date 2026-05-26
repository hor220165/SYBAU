from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)


OUT = Path("/Users/david/SYBAU/document_output/Projektabschlussbericht_SYBAU.docx")
PDF_OUT = Path("/Users/david/SYBAU/document_output/Projektabschlussbericht_SYBAU.pdf")

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(89, 89, 89)
LIGHT_FILL = "F2F4F7"
CALLOUT_FILL = "F4F6F9"
BORDER = "D9E2F3"


def de_text(text: str) -> str:
    replacements = {
        "Inhaltsuebersicht": "Inhaltsübersicht",
        "Projektueberblick": "Projektüberblick",
        "ueber": "über",
        "Ueber": "Über",
        "fuer": "für",
        "Fuer": "Für",
        "koennen": "können",
        "Koennen": "Können",
        "Uebungen": "Übungen",
        "Uebersicht": "Übersicht",
        "uebersicht": "übersicht",
        "funktionsfaehige": "funktionsfähige",
        "funktionsfaehig": "funktionsfähig",
        "oeffnen": "öffnen",
        "urspruenglichen": "ursprünglichen",
        "urspruengliche": "ursprüngliche",
        "urspruenglich": "ursprünglich",
        "vollstaendige": "vollständige",
        "vollstaendig": "vollständig",
        "erhoeht": "erhöht",
        "fuehrte": "führte",
        "fuehrung": "führung",
        "Nutzerfuehrung": "Nutzerführung",
        "herunterzaehlendem": "herunterzählendem",
        "herunterzaehlender": "herunterzählender",
        "herunterzaehlend": "herunterzählend",
        "zaehlendem": "zählendem",
        "Oeffentliche": "Öffentliche",
        "oeffentliche": "öffentliche",
        "Praesentationsbereit": "Präsentationsbereit",
        "Praesentationsqualitaet": "Präsentationsqualität",
        "Praesentation": "Präsentation",
        "Präsentationsqualitaet": "Präsentationsqualität",
        "praesentationsbereit": "präsentationsbereit",
        "praesentation": "präsentation",
        "praesentierbarer": "präsentierbarer",
        "praesentierbar": "präsentierbar",
        "praesentiert": "präsentiert",
        "Qualitaet": "Qualität",
        "Erfuellt": "Erfüllt",
        "Datenschutzerklaerung": "Datenschutzerklärung",
        "Regelmaessiges": "Regelmäßiges",
        "laeuft": "läuft",
        "zusaetzlich": "zusätzlich",
        "zusaetzlicher": "zusätzlicher",
        "zukuenftige": "zukünftige",
        "ermoeglichte": "ermöglichte",
        "ermoeglicht": "ermöglicht",
        "tatsaechlichen": "tatsächlichen",
        "Nutzerablaeufe": "Nutzerabläufe",
        "Einschraenkung": "Einschränkung",
        "Aktivitaetsdaten": "Aktivitätsdaten",
        "eingefuegt": "eingefügt",
        "geprueft": "geprüft",
        "geschuetzte": "geschützte",
        "waehrend": "während",
        "waere": "wäre",
        "frueh": "früh",
        "spaetere": "spätere",
        "beruecksichtigt": "berücksichtigt",
        "veraendert": "verändert",
        "groesser": "größer",
        "abschliessen": "abschließen",
        "abschliessen,": "abschließen,",
        "abschliessen.": "abschließen.",
    }
    for source, target in replacements.items():
        text = text.replace(source, target)
    return text


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_width(table, width_dxa: int = 9360) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")


def set_cell_margins(table, top=80, start=120, bottom=80, end=120) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_cell_mar = tbl_pr.find(qn("w:tblCellMar"))
    if tbl_cell_mar is None:
        tbl_cell_mar = OxmlElement("w:tblCellMar")
        tbl_pr.append(tbl_cell_mar)
    for key, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tbl_cell_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tbl_cell_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color: str = BORDER) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def style_run(run, bold=False, color=None, size=None) -> None:
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    run.bold = bold
    if color:
        run.font.color.rgb = color
    if size:
        run.font.size = Pt(size)


def add_para(doc, text: str = "", *, bold_prefix: str | None = None):
    text = de_text(text)
    bold_prefix = de_text(bold_prefix) if bold_prefix else None
    p = doc.add_paragraph()
    p.style = doc.styles["Normal"]
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        style_run(r, bold=True)
        r = p.add_run(text[len(bold_prefix):])
        style_run(r)
    else:
        r = p.add_run(text)
        style_run(r)
    return p


def add_bullets(doc, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.167
        run = p.add_run(de_text(item))
        style_run(run)


def add_numbered(doc, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.167
        run = p.add_run(de_text(item))
        style_run(run)


def add_callout(doc, title: str, body: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table)
    set_table_borders(table, "E1E7EF")
    set_cell_margins(table, top=120, bottom=120, start=160, end=160)
    cell = table.cell(0, 0)
    set_cell_shading(cell, CALLOUT_FILL)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(de_text(title))
    style_run(r, bold=True, color=DARK_BLUE)
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(de_text(body))
    style_run(r)
    doc.add_paragraph()


def add_kv_table(doc, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    set_table_width(table)
    set_table_borders(table)
    set_cell_margins(table)
    widths = [2160, 7200]
    for label, value in rows:
        row = table.add_row()
        for idx, text in enumerate((de_text(label), de_text(value))):
            cell = row.cells[idx]
            set_cell_width(cell, widths[idx])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if idx == 0:
                set_cell_shading(cell, LIGHT_FILL)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(text)
            style_run(run, bold=idx == 0, color=DARK_BLUE if idx == 0 else None)
    doc.add_paragraph()


def add_feature_table(doc, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    set_table_width(table)
    set_table_borders(table)
    set_cell_margins(table)
    widths = [2520, 6840]
    headers = ["Bereich", "Umgesetzte Funktion"]
    for idx, h in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_width(cell, widths[idx])
        set_cell_shading(cell, LIGHT_FILL)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(de_text(h))
        style_run(run, bold=True, color=DARK_BLUE)
    for area, desc in rows:
        row = table.add_row()
        for idx, text in enumerate((de_text(area), de_text(desc))):
            cell = row.cells[idx]
            set_cell_width(cell, widths[idx])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(text)
            style_run(run, bold=idx == 0, color=DARK_BLUE if idx == 0 else None)
    doc.add_paragraph()


def add_status_table(doc, rows: list[tuple[str, str, str]]) -> None:
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    set_table_width(table)
    set_table_borders(table)
    set_cell_margins(table)
    widths = [2200, 1540, 5620]
    headers = ["Bereich", "Status", "Kurzbewertung"]
    for idx, h in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_width(cell, widths[idx])
        set_cell_shading(cell, LIGHT_FILL)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(de_text(h))
        style_run(run, bold=True, color=DARK_BLUE)
    for area, status, desc in rows:
        row = table.add_row()
        for idx, text in enumerate((de_text(area), de_text(status), de_text(desc))):
            cell = row.cells[idx]
            set_cell_width(cell, widths[idx])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(text)
            style_run(run, bold=idx < 2, color=DARK_BLUE if idx == 0 else None)
    doc.add_paragraph()


def add_footer(section) -> None:
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.paragraph_format.space_before = Pt(0)
    footer.paragraph_format.space_after = Pt(0)
    run = footer.add_run("SYBAU Projektabschlussbericht")
    style_run(run, color=MUTED, size=9)


def configure_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    add_footer(section)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    title = styles["Title"]
    title.font.name = "Calibri"
    title._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    title.font.size = Pt(24)
    title.font.bold = True
    title.font.color.rgb = RGBColor(11, 37, 69)
    title.paragraph_format.space_after = Pt(6)

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def build() -> None:
    doc = Document()
    configure_doc(doc)

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("Projektabschlussbericht").bold = True

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("SYBAU - Shape Your Body And Unleash")
    style_run(r, bold=True, color=BLUE, size=15)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = meta.add_run("Team DidimDynamics | Stand: 26.05.2026")
    style_run(r, color=MUTED, size=10)

    add_kv_table(
        doc,
        [
            ("Projekt", "SYBAU - Shape Your Body And Unleash"),
            ("Team", "DidimDynamics"),
            ("Frontend", "Vue 3, TypeScript, Vite - Deployment ueber Vercel"),
            ("Backend", "ASP.NET Core / C# - Deployment ueber Render"),
            ("Datenbank", "Neon Free PostgreSQL"),
            ("Mobile App", "Flutter / Dart"),
            ("Oeffentliche Webadresse", "https://sybau-fitness.vercel.app"),
            ("Backend-URL", "https://sybau-xll5.onrender.com"),
        ],
    )

    add_callout(
        doc,
        "Hinweis zur Zeitaufzeichnung",
        "Die aktuelle Zeitaufzeichnung und die monatliche Auswertung werden separat eingefuegt. Dieser Bericht konzentriert sich auf Features, Soll-Ist-Vergleich, finalen Status und Retrospective.",
    )

    doc.add_heading(de_text("Inhaltsuebersicht"), level=1)
    add_numbered(
        doc,
        [
            "Zusammenfassung",
            "Projektueberblick",
            "Technische Umsetzung",
            "Umgesetzte Features",
            "Soll-Ist-Vergleich",
            "Finaler Statusreport",
            "Lessons Learned / Retrospective",
        ],
    )

    doc.add_page_break()

    doc.add_heading(de_text("1. Zusammenfassung"), level=1)
    add_para(
        doc,
        "Das Projekt SYBAU wurde am Projektende als funktionsfaehige Fitness-Plattform mit Website, Backend, Datenbank und Mobile-App fertiggestellt. Die Kernidee, Training durch Gamification sichtbarer und motivierender zu machen, wurde umgesetzt. Nutzer koennen Uebungen eintragen, XP und Coins sammeln, Quests abschliessen, Achievements freischalten, Chests oeffnen und ihren Avatar weiterentwickeln.",
    )
    add_para(
        doc,
        "Gegenueber dem urspruenglichen Projektstand wurde der Umfang erweitert. Besonders wichtig ist, dass neben der Web-Anwendung auch eine vollstaendige Mobile-App umgesetzt wurde. Diese Mobile-App war im urspruenglichen Umfang nicht als Hauptliefergegenstand geplant, wurde aber als Zusatzumfang realisiert und erhoeht den praktischen Nutzen des Projekts deutlich.",
    )
    add_para(
        doc,
        "Der finale Stand ist fuer die Projektabschlusspräsentation geeignet. Die wichtigsten Funktionen sind implementiert und die Anwendung kann anhand realer Nutzerablaeufe praesentiert werden.",
    )

    doc.add_heading(de_text("2. Projektueberblick"), level=1)
    add_para(
        doc,
        "SYBAU ist eine digitale Fitness-Anwendung, die klassisches Workout-Tracking mit spielerischen Elementen verbindet. Ziel war es, Fortschritt nicht nur als Zahlenliste darzustellen, sondern als sichtbare Entwicklung ueber Level, Avatar, Belohnungen und soziale Vergleiche.",
    )
    add_para(
        doc,
        "Die Plattform besteht aus einer Website, einer Mobile-App und einem gemeinsamen Backend. Website und Mobile-App greifen auf dieselbe API und dieselbe Datenbank zu. Dadurch bleiben Konto, Fortschritt, Quests, Shop-Inhalte, Profil und Leaderboard auf beiden Plattformen konsistent.",
    )
    add_bullets(
        doc,
        [
            "Zielgruppe: Nutzer, die Training strukturierter und motivierender verfolgen wollen.",
            "Kernnutzen: schneller Trainingseintrag, sichtbarer Fortschritt und Belohnungssystem.",
            "Motivationsansatz: XP, Coins, Quests, Achievements, Avatar-Entwicklung und Leaderboard.",
            "Ergebnis: praesentierbarer Prototyp mit Web-App, Mobile-App und produktionsnaher Infrastruktur.",
        ],
    )

    doc.add_heading(de_text("3. Technische Umsetzung"), level=1)
    add_para(
        doc,
        "Die technische Umsetzung wurde in drei Hauptbereiche aufgeteilt: Web-Frontend, Mobile-App und Backend. Das Backend stellt die zentrale Schnittstelle bereit und verwaltet Authentifizierung, Nutzerfortschritt, Workouts, Quests, Shop, Chests, Achievements, Freunde und Leaderboard.",
    )
    add_feature_table(
        doc,
        [
            ("Web-Frontend", "Vue 3, TypeScript und Vite. Die Website wird ueber Vercel bereitgestellt und ist unter https://sybau-fitness.vercel.app erreichbar."),
            ("Mobile-App", "Flutter und Dart. Die App nutzt dieselbe API wie die Website und bildet die wichtigsten Funktionen fuer mobile Nutzung ab."),
            ("Backend", "ASP.NET Core / C# mit Entity Framework Core. Das Backend laeuft auf Render unter https://sybau-xll5.onrender.com."),
            ("Datenbank", "Neon Free PostgreSQL in Produktion. Lokal kann fuer Entwicklung SQLite verwendet werden."),
            ("Sicherheit", "JWT-basierte Authentifizierung, Admin-Schutz, Eingabevalidierung bei Login/Registrierung und Datenschutz-/Impressum-Seiten."),
            ("Deployment", "Frontend ueber Vercel, Backend ueber Render, Datenbank ueber Neon Free PostgreSQL."),
        ],
    )

    doc.add_heading(de_text("4. Umgesetzte Features"), level=1)
    doc.add_heading(de_text("4.1 Website"), level=2)
    add_feature_table(
        doc,
        [
            ("Landingpage", "Oeffentliche Startseite mit Projektvorstellung, Call-to-Action zur Registrierung und responsiver Darstellung."),
            ("Auth", "Registrierung, Login, E-Mail-Pruefung, Passwort-Hinweise, Datenschutz-Checkbox und Schutz gegen Temp-Mail-Domains."),
            ("Dashboard", "Zentrale Uebersicht mit Avatar, Level, XP, Coins, Streaks und direktem Fortschrittsfeedback."),
            ("Workouts", "50 Uebungen mit Kategorien, Suchfeld, Schwierigkeit, Eingabeformat Zeit oder Einheiten sowie balancierten XP-/Coin-Belohnungen."),
            ("Quests", "Taegliche, woechentliche und monatliche Quests mit Fortschritt, Belohnungen und live herunterzaehlendem Daily-Timer."),
            ("Shop", "Items, Booster, Chests, Coin-Pakete und Chest-Opening. Echtgeld-Angebote sind vorbereitet, aber als 'noch in Arbeit' gekennzeichnet."),
            ("Avatar", "Avatar-System mit sichtbarer Entwicklung und Item-/Booster-Bezug."),
            ("Profil", "Profilbild, Benutzername, privates Profil, Achievements, Statistiken und Weekly Activity Kalender."),
            ("Freunde", "Nutzersuche ab sinnvoller Zeichenzahl, Freundschaftsanfragen, oeffentliche Profile und soziale Vergleiche."),
            ("Leaderboard", "Globale Rangliste mit Level- und XP-Vergleich."),
            ("Admin", "Admin-Bereich zur Verwaltung von Inhalten, unter anderem Shop-Items und Item-Typen."),
            ("Rechtliches", "Impressum und Datenschutzerklaerung mit Scroll-to-Top-Verhalten und klarer Beschreibung der genutzten Daten."),
        ],
    )

    doc.add_heading(de_text("4.2 Mobile-App"), level=2)
    add_para(
        doc,
        "Die Mobile-App wurde als Zusatzumfang umgesetzt. Sie war urspruenglich nicht als Hauptbestandteil des Projekts geplant, wurde aber fertiggestellt, um SYBAU auch unterwegs sinnvoll nutzbar zu machen.",
    )
    add_feature_table(
        doc,
        [
            ("Login und Konto", "Nutzer koennen sich mit demselben Konto wie auf der Website anmelden."),
            ("Dashboard", "Mobile Uebersicht mit Avatar, Level, XP, Coins und Fortschritt."),
            ("Workouts", "Mobile Workout-Erfassung mit Suche, Kategorien und schneller Bedienung waehrend oder nach dem Training."),
            ("Quests", "Taegliche, woechentliche und monatliche Quests inklusive live herunterzaehlendem Daily-Timer."),
            ("Profil", "Profilansicht mit Achievements, Aktivitaetskalender, Profilbild und Sichtbarkeitseinstellungen."),
            ("Freunde und Leaderboard", "Soziale Funktionen wie Freunde, oeffentliche Profile und Rangliste sind mobil nutzbar."),
            ("Health-Sync", "Mobile Aktivitaetsdaten koennen fuer Fortschritt und Quest-Fortschritt genutzt werden bzw. sind dafuer vorbereitet."),
        ],
    )

    doc.add_heading(de_text("5. Soll-Ist-Vergleich"), level=1)
    add_status_table(
        doc,
        [
            ("Account und Auth", "Erfuellt", "Registrierung, Login, Validierung, Datenschutz-Checkbox und geschuetzte Bereiche wurden umgesetzt."),
            ("Workout-System", "Erfuellt", "Die geplante Trainingserfassung wurde mit 50 Uebungen, Suche, Kategorien und balancierten Belohnungen erweitert."),
            ("XP, Coins und Level", "Erfuellt", "Fortschritt wird berechnet, gespeichert und direkt im Header sichtbar gemacht."),
            ("Avatar-System", "Erfuellt", "Der Avatar ist als sichtbares Fortschrittselement integriert und mit dem Level-/Item-System verbunden."),
            ("Quests", "Erfuellt", "Daily, Weekly und Monthly Quests sind umgesetzt; Daily besitzt einen live laufenden Timer."),
            ("Shop und Chests", "Erfuellt", "Shop, Items, Booster und Chest-Opening sind umgesetzt. Echtgeldkaeufe sind vorbereitet, aber ohne Zahlungssystem."),
            ("Freunde und Leaderboard", "Erfuellt", "Freundesfunktionen, oeffentliche Profile und Ranglisten wurden umgesetzt."),
            ("Deployment", "Erfuellt", "Frontend, Backend und Datenbank sind produktionsnah ueber Vercel, Render und Neon Free erreichbar."),
            ("Mobile-App", "Zusatzumfang", "Eine vollstaendige Flutter-App wurde zusaetzlich fertiggestellt und ist damit eine Scope-Erweiterung gegenueber dem urspruenglichen Web-Fokus."),
            ("Google-Fit/Web-Anbindung", "Angepasst", "Eine direkte Google-Fit-Anbindung in der Web-App wurde nicht als Kernfunktion umgesetzt. Stattdessen wurde der mobile Health-Sync-Ansatz verfolgt."),
            ("Echtgeld-Zahlung", "Vorbereitet", "Echtgeld-Angebote koennen dargestellt werden. Die echte Zahlungsabwicklung wurde bewusst nicht implementiert und zeigt aktuell 'noch in Arbeit'."),
        ],
    )
    add_para(
        doc,
        "Insgesamt wurde der geplante Kernumfang erreicht. Einzelne Punkte wurden im Projektverlauf angepasst, weil sich technische oder zeitliche Rahmenbedingungen veraendert haben. Gleichzeitig wurde mit der Mobile-App ein zusaetzlicher Liefergegenstand fertiggestellt, der den urspruenglichen Nutzen des Projekts erweitert.",
    )

    doc.add_heading(de_text("6. Finaler Statusreport"), level=1)
    add_status_table(
        doc,
        [
            ("Web-Anwendung", "Fertig", "Die Website ist funktionsfaehig, responsiv ueberarbeitet und fuer die Live-Demonstration geeignet."),
            ("Mobile-App", "Fertig", "Die Flutter-App deckt die wichtigsten Nutzerablaeufe ab und nutzt dieselbe Backend-API."),
            ("Backend", "Fertig", "Die API stellt die zentralen Funktionen bereit und baut erfolgreich."),
            ("Datenbank", "Fertig", "Die Produktionsdaten werden in Neon Free PostgreSQL gespeichert."),
            ("Rechtliches", "Fertig", "Datenschutz und Impressum sind vorhanden und an die tatsaechlich genutzten Dienste angepasst."),
            ("Qualitaet", "Praesentationsbereit", "Web-Build, Backend-Build, Flutter-Tests, Flutter-Analyse und iOS-Simulator-Build wurden geprueft."),
            ("Bekannte Einschraenkung", "Akzeptiert", "Echte Zahlungsabwicklung fuer Echtgeld-Angebote ist nicht Teil der finalen Implementierung."),
        ],
    )
    add_callout(
        doc,
        "Finale Bewertung",
        "SYBAU ist als Projektabschluss praesentationsbereit. Die Kernfunktionen sind umgesetzt, der Zusatzumfang Mobile-App ist fertiggestellt und die technische Infrastruktur ist fuer eine Live-Produktpraesentation ausreichend stabil.",
    )

    doc.add_heading(de_text("7. Lessons Learned / Retrospective"), level=1)
    doc.add_heading(de_text("Was gut funktioniert hat"), level=2)
    add_bullets(
        doc,
        [
            "Die Aufteilung in Frontend, Backend und Mobile-App hat geholfen, parallel an mehreren Bereichen zu arbeiten.",
            "Gamification-Elemente wie XP, Coins, Quests und Chests haben das Projekt klarer und motivierender gemacht.",
            "Regelmaessiges Feedback fuehrte zu vielen konkreten Verbesserungen bei Responsiveness, Datenschutz, Navigation und Nutzerfuehrung.",
            "Die gemeinsame API fuer Website und Mobile-App war sinnvoll, weil beide Plattformen dieselben Daten verwenden koennen.",
        ],
    )
    doc.add_heading(de_text("Was herausfordernd war"), level=2)
    add_bullets(
        doc,
        [
            "Die Abstimmung zwischen Web, Mobile, Backend und Datenbank wurde mit wachsendem Funktionsumfang komplexer.",
            "Deployment und Hosting mussten mehrfach angepasst werden, bis Frontend, Backend und Datenbank gut zusammenarbeiteten.",
            "Responsive Darstellung war besonders wichtig, da Website und Mobile-App beide sauber nutzbar sein sollten.",
            "Ein echtes Zahlungssystem waere fuer den Projektumfang zu gross gewesen und wurde daher bewusst nur vorbereitet.",
        ],
    )
    doc.add_heading(de_text("Erkenntnisse fuer zukuenftige Projekte"), level=2)
    add_bullets(
        doc,
        [
            "Technische Infrastruktur sollte frueh verbindlich festgelegt werden, damit spaetere Hosting-Wechsel weniger Aufwand verursachen.",
            "Mobile Anforderungen sollten frueh in der Planung beruecksichtigt werden, auch wenn sie zuerst nur als Zusatzidee wirken.",
            "Rechtliche Texte und Datenschutz sollten nicht erst am Ende entstehen, sondern parallel zur tatsaechlichen Datennutzung gepflegt werden.",
            "Kleine UI-Details wie Ladezustand, Tooltips, Timer und responsives Verhalten haben grossen Einfluss auf die Praesentationsqualitaet.",
        ],
    )

    doc.add_heading(de_text("Abschluss"), level=1)
    add_para(
        doc,
        "Das Projektziel wurde erreicht. SYBAU liegt als funktionsfaehiges, praesentierbares System mit Website, Backend, Datenbank und Mobile-App vor. Besonders positiv ist, dass der Projektumfang ueber den urspruenglichen Web-Fokus hinaus erweitert wurde und mit der Mobile-App ein zusaetzlicher praktischer Zugang zum System entstanden ist.",
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    build_pdf()


def pdf_styles():
    base = getSampleStyleSheet()
    styles = {
        "title": ParagraphStyle(
            "SybauTitle",
            parent=base["Title"],
            fontName="Helvetica-Bold",
            fontSize=24,
            textColor=colors.HexColor("#0B2545"),
            alignment=TA_CENTER,
            spaceAfter=8,
        ),
        "subtitle": ParagraphStyle(
            "SybauSubtitle",
            parent=base["Normal"],
            fontName="Helvetica-Bold",
            fontSize=14,
            textColor=colors.HexColor("#2E74B5"),
            alignment=TA_CENTER,
            spaceAfter=14,
        ),
        "meta": ParagraphStyle(
            "SybauMeta",
            parent=base["Normal"],
            fontName="Helvetica",
            fontSize=9,
            textColor=colors.HexColor("#595959"),
            alignment=TA_CENTER,
            spaceAfter=14,
        ),
        "h1": ParagraphStyle(
            "SybauH1",
            parent=base["Heading1"],
            fontName="Helvetica-Bold",
            fontSize=15,
            textColor=colors.HexColor("#2E74B5"),
            spaceBefore=14,
            spaceAfter=7,
        ),
        "h2": ParagraphStyle(
            "SybauH2",
            parent=base["Heading2"],
            fontName="Helvetica-Bold",
            fontSize=12,
            textColor=colors.HexColor("#2E74B5"),
            spaceBefore=10,
            spaceAfter=5,
        ),
        "body": ParagraphStyle(
            "SybauBody",
            parent=base["BodyText"],
            fontName="Helvetica",
            fontSize=9.5,
            leading=12.2,
            spaceAfter=6,
        ),
        "bullet": ParagraphStyle(
            "SybauBullet",
            parent=base["BodyText"],
            fontName="Helvetica",
            fontSize=9.5,
            leading=12.2,
            leftIndent=14,
            firstLineIndent=-8,
            spaceAfter=3,
        ),
        "small": ParagraphStyle(
            "SybauSmall",
            parent=base["BodyText"],
            fontName="Helvetica",
            fontSize=8.5,
            leading=10.5,
            spaceAfter=0,
        ),
        "table_head": ParagraphStyle(
            "SybauTableHead",
            parent=base["BodyText"],
            fontName="Helvetica-Bold",
            fontSize=8.8,
            textColor=colors.HexColor("#1F4D78"),
            leading=10.5,
        ),
        "table": ParagraphStyle(
            "SybauTable",
            parent=base["BodyText"],
            fontName="Helvetica",
            fontSize=8.4,
            leading=10.2,
        ),
        "table_bold": ParagraphStyle(
            "SybauTableBold",
            parent=base["BodyText"],
            fontName="Helvetica-Bold",
            fontSize=8.4,
            textColor=colors.HexColor("#1F4D78"),
            leading=10.2,
        ),
    }
    return styles


def p(text: str, style):
    return Paragraph(de_text(text).replace("&", "&amp;"), style)


def pdf_table(data, widths, styles, header=True):
    rows = []
    for row_i, row in enumerate(data):
        pdf_row = []
        for col_i, text in enumerate(row):
            style = styles["table_head"] if header and row_i == 0 else styles["table_bold"] if col_i == 0 else styles["table"]
            pdf_row.append(p(text, style))
        rows.append(pdf_row)
    table = Table(rows, colWidths=widths, hAlign="LEFT", repeatRows=1 if header else 0)
    commands = [
        ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#D9E2F3")),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 6),
        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
        ("TOPPADDING", (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
    ]
    if header:
        commands.append(("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#F2F4F7")))
    table.setStyle(TableStyle(commands))
    return table


def pdf_para(story, text, styles):
    story.append(p(text, styles["body"]))


def pdf_bullets(story, items, styles):
    for item in items:
        story.append(p(f"- {item}", styles["bullet"]))


def add_pdf_footer(canvas, doc):
    canvas.saveState()
    canvas.setFont("Helvetica", 8)
    canvas.setFillColor(colors.HexColor("#777777"))
    canvas.drawRightString(7.5 * inch, 0.55 * inch, f"SYBAU Projektabschlussbericht | Seite {doc.page}")
    canvas.restoreState()


def build_pdf() -> None:
    styles = pdf_styles()
    doc = SimpleDocTemplate(
        str(PDF_OUT),
        pagesize=letter,
        rightMargin=0.75 * inch,
        leftMargin=0.75 * inch,
        topMargin=0.75 * inch,
        bottomMargin=0.75 * inch,
    )
    story = []

    story.append(p("Projektabschlussbericht", styles["title"]))
    story.append(p("SYBAU - Shape Your Body And Unleash", styles["subtitle"]))
    story.append(p("Team DidimDynamics | Stand: 26.05.2026", styles["meta"]))
    story.append(
        pdf_table(
            [
                ("Projekt", "SYBAU - Shape Your Body And Unleash"),
                ("Team", "DidimDynamics"),
                ("Frontend", "Vue 3, TypeScript, Vite - Deployment ueber Vercel"),
                ("Backend", "ASP.NET Core / C# - Deployment ueber Render"),
                ("Datenbank", "Neon Free PostgreSQL"),
                ("Mobile App", "Flutter / Dart"),
                ("Oeffentliche Webadresse", "https://sybau-fitness.vercel.app"),
                ("Backend-URL", "https://sybau-xll5.onrender.com"),
            ],
            [1.65 * inch, 5.25 * inch],
            styles,
            header=False,
        )
    )
    story.append(Spacer(1, 12))
    story.append(
        pdf_table(
            [
                (
                    "Hinweis",
                    "Die aktuelle Zeitaufzeichnung und die monatliche Auswertung werden separat eingefuegt. Dieser Bericht konzentriert sich auf Features, Soll-Ist-Vergleich, finalen Status und Retrospective.",
                )
            ],
            [1.0 * inch, 5.9 * inch],
            styles,
            header=False,
        )
    )
    story.append(PageBreak())

    story.append(p("1. Zusammenfassung", styles["h1"]))
    pdf_para(story, "Das Projekt SYBAU wurde am Projektende als funktionsfaehige Fitness-Plattform mit Website, Backend, Datenbank und Mobile-App fertiggestellt. Die Kernidee, Training durch Gamification sichtbarer und motivierender zu machen, wurde umgesetzt. Nutzer koennen Uebungen eintragen, XP und Coins sammeln, Quests abschliessen, Achievements freischalten, Chests oeffnen und ihren Avatar weiterentwickeln.", styles)
    pdf_para(story, "Gegenueber dem urspruenglichen Projektstand wurde der Umfang erweitert. Besonders wichtig ist, dass neben der Web-Anwendung auch eine vollstaendige Mobile-App umgesetzt wurde. Diese Mobile-App war im urspruenglichen Umfang nicht als Hauptliefergegenstand geplant, wurde aber als Zusatzumfang realisiert und erhoeht den praktischen Nutzen des Projekts deutlich.", styles)
    pdf_para(story, "Der finale Stand ist fuer die Projektabschlusspräsentation geeignet. Die wichtigsten Funktionen sind implementiert und die Anwendung kann anhand realer Nutzerablaeufe praesentiert werden.", styles)

    story.append(p("2. Projektueberblick", styles["h1"]))
    pdf_para(story, "SYBAU ist eine digitale Fitness-Anwendung, die klassisches Workout-Tracking mit spielerischen Elementen verbindet. Ziel war es, Fortschritt nicht nur als Zahlenliste darzustellen, sondern als sichtbare Entwicklung ueber Level, Avatar, Belohnungen und soziale Vergleiche.", styles)
    pdf_para(story, "Die Plattform besteht aus einer Website, einer Mobile-App und einem gemeinsamen Backend. Website und Mobile-App greifen auf dieselbe API und dieselbe Datenbank zu. Dadurch bleiben Konto, Fortschritt, Quests, Shop-Inhalte, Profil und Leaderboard auf beiden Plattformen konsistent.", styles)
    pdf_bullets(story, [
        "Zielgruppe: Nutzer, die Training strukturierter und motivierender verfolgen wollen.",
        "Kernnutzen: schneller Trainingseintrag, sichtbarer Fortschritt und Belohnungssystem.",
        "Motivationsansatz: XP, Coins, Quests, Achievements, Avatar-Entwicklung und Leaderboard.",
        "Ergebnis: praesentierbarer Prototyp mit Web-App, Mobile-App und produktionsnaher Infrastruktur.",
    ], styles)

    story.append(p("3. Technische Umsetzung", styles["h1"]))
    pdf_para(story, "Die technische Umsetzung wurde in drei Hauptbereiche aufgeteilt: Web-Frontend, Mobile-App und Backend. Das Backend stellt die zentrale Schnittstelle bereit und verwaltet Authentifizierung, Nutzerfortschritt, Workouts, Quests, Shop, Chests, Achievements, Freunde und Leaderboard.", styles)
    story.append(pdf_table([
        ("Bereich", "Umsetzung"),
        ("Web-Frontend", "Vue 3, TypeScript und Vite. Deployment ueber Vercel unter https://sybau-fitness.vercel.app."),
        ("Mobile-App", "Flutter und Dart. Die App nutzt dieselbe API wie die Website."),
        ("Backend", "ASP.NET Core / C# mit Entity Framework Core. Betrieb ueber Render unter https://sybau-xll5.onrender.com."),
        ("Datenbank", "Neon Free PostgreSQL in Produktion. SQLite fuer lokale Entwicklung."),
        ("Sicherheit", "JWT-Authentifizierung, Admin-Schutz, Eingabevalidierung und Datenschutz-/Impressum-Seiten."),
    ], [1.55 * inch, 5.35 * inch], styles))

    story.append(p("4. Umgesetzte Features", styles["h1"]))
    story.append(p("4.1 Website", styles["h2"]))
    story.append(pdf_table([
        ("Bereich", "Umgesetzte Funktion"),
        ("Landingpage", "Projektvorstellung, Call-to-Action zur Registrierung und responsive Darstellung."),
        ("Auth", "Registrierung, Login, E-Mail-Pruefung, Passwort-Hinweise, Datenschutz-Checkbox und Temp-Mail-Schutz."),
        ("Dashboard", "Uebersicht mit Avatar, Level, XP, Coins, Streaks und Fortschrittsfeedback."),
        ("Workouts", "50 Uebungen mit Kategorien, Suche, Schwierigkeit, Zeit/Einheiten und balancierten Belohnungen."),
        ("Quests", "Daily, Weekly und Monthly Quests mit Fortschritt, Rewards und live Daily-Timer."),
        ("Shop", "Items, Booster, Chests, Coin-Pakete und Chest-Opening. Echtgeld-Angebote sind vorbereitet."),
        ("Profil", "Profilbild, Name, privates Profil, Achievements, Statistiken und Weekly Activity."),
        ("Freunde/Leaderboard", "Nutzersuche, Freundschaftsanfragen, oeffentliche Profile und Rangliste."),
        ("Admin/Rechtliches", "Admin-Bereich, Impressum und Datenschutzerklaerung."),
    ], [1.55 * inch, 5.35 * inch], styles))

    story.append(p("4.2 Mobile-App", styles["h2"]))
    pdf_para(story, "Die Mobile-App wurde als Zusatzumfang umgesetzt. Sie war urspruenglich nicht als Hauptbestandteil des Projekts geplant, wurde aber fertiggestellt, um SYBAU auch unterwegs sinnvoll nutzbar zu machen.", styles)
    story.append(pdf_table([
        ("Bereich", "Umgesetzte Funktion"),
        ("Login und Konto", "Nutzung desselben Accounts wie auf der Website."),
        ("Dashboard", "Mobile Uebersicht mit Avatar, Level, XP, Coins und Fortschritt."),
        ("Workouts", "Mobile Workout-Erfassung mit Suche, Kategorien und schneller Bedienung."),
        ("Quests", "Daily, Weekly und Monthly Quests inklusive live Daily-Timer."),
        ("Profil", "Achievements, Aktivitaetskalender, Profilbild und Sichtbarkeitseinstellungen."),
        ("Soziale Funktionen", "Freunde, oeffentliche Profile und Leaderboard."),
        ("Health-Sync", "Mobile Aktivitaetsdaten koennen fuer Fortschritt und Quests genutzt werden bzw. sind vorbereitet."),
    ], [1.55 * inch, 5.35 * inch], styles))

    story.append(p("5. Soll-Ist-Vergleich", styles["h1"]))
    story.append(pdf_table([
        ("Bereich", "Status", "Kurzbewertung"),
        ("Account und Auth", "Erfuellt", "Login, Registrierung, Validierung, Datenschutz-Checkbox und geschuetzte Bereiche wurden umgesetzt."),
        ("Workout-System", "Erfuellt", "Trainingserfassung wurde mit 50 Uebungen, Suche, Kategorien und Rewards erweitert."),
        ("XP, Coins und Level", "Erfuellt", "Fortschritt wird berechnet, gespeichert und direkt sichtbar gemacht."),
        ("Avatar-System", "Erfuellt", "Avatar ist als Fortschrittselement integriert."),
        ("Quests", "Erfuellt", "Daily, Weekly und Monthly Quests inklusive live Daily-Timer."),
        ("Shop und Chests", "Erfuellt", "Shop, Items, Booster und Chest-Opening. Echtgeld ohne Zahlungssystem vorbereitet."),
        ("Mobile-App", "Zusatzumfang", "Vollstaendige Flutter-App als Scope-Erweiterung gegenueber dem urspruenglichen Web-Fokus."),
        ("Google-Fit/Web", "Angepasst", "Direkte Web-Anbindung wurde durch mobilen Health-Sync-Ansatz ersetzt."),
        ("Deployment", "Erfuellt", "Vercel, Render und Neon Free sind eingerichtet."),
    ], [1.45 * inch, 1.05 * inch, 4.4 * inch], styles))
    pdf_para(story, "Insgesamt wurde der geplante Kernumfang erreicht. Einzelne Punkte wurden angepasst, weil sich technische oder zeitliche Rahmenbedingungen veraendert haben. Gleichzeitig wurde mit der Mobile-App ein zusaetzlicher Liefergegenstand fertiggestellt.", styles)

    story.append(p("6. Finaler Statusreport", styles["h1"]))
    story.append(pdf_table([
        ("Bereich", "Status", "Kurzbewertung"),
        ("Web-Anwendung", "Fertig", "Funktionsfaehig, responsiv ueberarbeitet und fuer Live-Demo geeignet."),
        ("Mobile-App", "Fertig", "Deckt die wichtigsten Nutzerablaeufe ab und nutzt dieselbe Backend-API."),
        ("Backend", "Fertig", "Zentrale API-Funktionen vorhanden und Build erfolgreich."),
        ("Datenbank", "Fertig", "Produktionsdaten in Neon Free PostgreSQL."),
        ("Qualitaet", "Praesentationsbereit", "Web-Build, Backend-Build, Flutter-Tests, Flutter-Analyse und iOS-Simulator-Build wurden geprueft."),
        ("Einschraenkung", "Akzeptiert", "Echte Zahlungsabwicklung fuer Echtgeld-Angebote ist nicht Teil der finalen Implementierung."),
    ], [1.45 * inch, 1.05 * inch, 4.4 * inch], styles))

    story.append(p("7. Lessons Learned / Retrospective", styles["h1"]))
    story.append(p("Was gut funktioniert hat", styles["h2"]))
    pdf_bullets(story, [
        "Die Aufteilung in Frontend, Backend und Mobile-App ermoeglichte paralleles Arbeiten.",
        "Gamification-Elemente machten das Projekt klarer und motivierender.",
        "Regelmaessiges Feedback fuehrte zu konkreten Verbesserungen bei Responsiveness, Datenschutz und Nutzerfuehrung.",
        "Die gemeinsame API war sinnvoll, weil Website und Mobile-App dieselben Daten verwenden.",
    ], styles)
    story.append(p("Was herausfordernd war", styles["h2"]))
    pdf_bullets(story, [
        "Die Abstimmung zwischen Web, Mobile, Backend und Datenbank wurde mit wachsendem Funktionsumfang komplexer.",
        "Deployment und Hosting mussten angepasst werden, bis Frontend, Backend und Datenbank stabil zusammenspielten.",
        "Responsive Darstellung war ein wichtiger Feinschliff fuer die Praesentationsqualitaet.",
        "Ein echtes Zahlungssystem waere fuer den Projektumfang zu gross gewesen und wurde daher bewusst vorbereitet, aber nicht abgeschlossen.",
    ], styles)
    story.append(p("Erkenntnisse fuer zukuenftige Projekte", styles["h2"]))
    pdf_bullets(story, [
        "Technische Infrastruktur sollte frueh verbindlich festgelegt werden.",
        "Mobile Anforderungen sollten frueher beruecksichtigt werden, auch wenn sie zuerst nur Zusatzidee sind.",
        "Datenschutztexte sollten parallel zur tatsaechlichen Datennutzung gepflegt werden.",
        "Kleine UI-Details wie Timer, Tooltips und Ladezustaende haben grossen Einfluss auf die Wirkung.",
    ], styles)

    story.append(p("Abschluss", styles["h1"]))
    pdf_para(story, "Das Projektziel wurde erreicht. SYBAU liegt als funktionsfaehiges, praesentierbares System mit Website, Backend, Datenbank und Mobile-App vor. Besonders positiv ist, dass der Projektumfang ueber den urspruenglichen Web-Fokus hinaus erweitert wurde und mit der Mobile-App ein zusaetzlicher praktischer Zugang zum System entstanden ist.", styles)

    doc.build(story, onFirstPage=add_pdf_footer, onLaterPages=add_pdf_footer)


if __name__ == "__main__":
    build()
