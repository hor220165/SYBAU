from __future__ import annotations

from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


TEMPLATE = Path("/Users/david/Desktop/4BHIF/PRE/Projektstatusbericht_DidimDynamics.docx")
OUT = Path("/Users/david/SYBAU/document_output/Projektabschlussbericht_SYBAU_Vorlage.docx")


def paragraph_text(element) -> str:
    return "".join(node.text or "" for node in element.xpath(".//w:t"))


def set_paragraph_text(paragraph, text: str) -> None:
    if not paragraph.runs:
        paragraph.add_run(text)
        return

    first = paragraph.runs[0]
    first.text = text
    for run in paragraph.runs[1:]:
        run.text = ""


def remove_old_report_content(doc: Document) -> None:
    body = doc._body._element
    children = list(body)
    start_index = None

    for index, child in enumerate(children):
        if child.tag == qn("w:p") and paragraph_text(child).strip() == "Inhaltsverzeichnis":
            start_index = index
            break

    if start_index is None:
        raise RuntimeError("Inhaltsverzeichnis marker not found in template.")

    for child in children[start_index:]:
        if child.tag == qn("w:sectPr"):
            continue
        body.remove(child)


def update_cover(doc: Document) -> None:
    replacements = {
        "Projektstatusbericht": "Projektabschlussbericht",
        "Berichtszeitraum: [17.2.2026 – 18.03.2026]": "Berichtszeitraum: [17.02.2026 – 26.05.2026]",
        "Datum: 18.0.2026": "Datum: 26.05.2026",
    }
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text in replacements:
            set_paragraph_text(paragraph, replacements[text])


def update_footer_text(doc: Document) -> None:
    for section in doc.sections:
        for paragraph in section.footer.paragraphs:
            for index, run in enumerate(paragraph.runs):
                if run.text == "Projektstatusbericht: SYBAU":
                    run.text = "Projektabschlussbericht: SYBAU"
                if index == 11 and run.text == "18":
                    run.text = "26"
                if index == 13 and run.text == "3":
                    run.text = "05"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_width(table, width_dxa: int = 9360) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")


def set_cell_margins(table, top=90, start=120, bottom=90, end=120) -> None:
    tbl_pr = table._tbl.tblPr
    margins = tbl_pr.find(qn("w:tblCellMar"))
    if margins is None:
        margins = OxmlElement("w:tblCellMar")
        tbl_pr.append(margins)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def style_table(table, widths: list[int], header: bool = True) -> None:
    table.style = "Table Grid"
    set_table_width(table)
    set_cell_margins(table)
    for row_index, row in enumerate(table.rows):
        for cell_index, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            set_cell_width(cell, widths[cell_index])
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.size = Pt(10)
            if header and row_index == 0:
                set_cell_shading(cell, "D9EAD3")
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True


def add_para(doc: Document, text: str, style: str | None = None) -> None:
    paragraph = doc.add_paragraph(style=style)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.add_run(text)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        paragraph = doc.add_paragraph(style="List Paragraph")
        paragraph.paragraph_format.space_after = Pt(3)
        paragraph.add_run(f"• {item}")


def add_numbered(doc: Document, items: list[str]) -> None:
    for index, item in enumerate(items, start=1):
        paragraph = doc.add_paragraph(style="List Paragraph")
        paragraph.paragraph_format.space_after = Pt(3)
        paragraph.add_run(f"{index}. {item}")


def add_feature_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Bereich"
    table.rows[0].cells[1].text = "Umgesetzte Funktion"
    for area, description in rows:
        cells = table.add_row().cells
        cells[0].text = area
        cells[1].text = description
    style_table(table, [2300, 7060])
    doc.add_paragraph()


def add_status_table(doc: Document, rows: list[tuple[str, str, str]]) -> None:
    table = doc.add_table(rows=1, cols=3)
    headers = ("Bereich", "Status", "Kurzbewertung")
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header
    for area, status, description in rows:
        cells = table.add_row().cells
        cells[0].text = area
        cells[1].text = status
        cells[2].text = description
    style_table(table, [1900, 1600, 5860])
    doc.add_paragraph()


def add_callout(doc: Document, title: str, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    cell.text = ""
    p = cell.paragraphs[0]
    title_run = p.add_run(title)
    title_run.bold = True
    p.add_run(f"\n{text}")
    set_cell_shading(cell, "EAF2F8")
    style_table(table, [9360], header=False)
    doc.add_paragraph()


def add_report_content(doc: Document) -> None:
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    section.header.is_linked_to_previous = True
    section.footer.is_linked_to_previous = True

    doc.add_heading("Inhaltsverzeichnis", level=1)
    add_numbered(
        doc,
        [
            "Zusammenfassung",
            "Projektüberblick",
            "Technische Umsetzung",
            "Umgesetzte Features",
            "Soll-Ist-Vergleich",
            "Finaler Statusreport",
            "Lessons Learned / Retrospective",
        ],
    )

    doc.add_heading("1. Zusammenfassung", level=1)
    add_para(
        doc,
        "Das Projekt SYBAU wurde am Projektende als funktionsfähige Fitness-Plattform mit Website, Backend, Datenbank und Mobile-App fertiggestellt. Die Kernidee, Training durch Gamification sichtbarer und motivierender zu machen, wurde umgesetzt. Nutzer können Übungen eintragen, XP und Coins sammeln, Quests abschließen, Achievements freischalten, Chests öffnen und ihren Avatar weiterentwickeln.",
    )
    add_para(
        doc,
        "Gegenüber dem ursprünglichen Projektstand wurde der Umfang erweitert. Besonders wichtig ist, dass neben der Web-Anwendung auch eine vollständige Mobile-App umgesetzt wurde. Diese Mobile-App war im ursprünglichen Umfang nicht als Hauptliefergegenstand geplant, wurde aber als Zusatzumfang realisiert und erhöht den praktischen Nutzen des Projekts deutlich.",
    )
    add_para(
        doc,
        "Der finale Stand ist für die Projektabschlusspräsentation geeignet. Die wichtigsten Funktionen sind implementiert und die Anwendung kann anhand realer Nutzerabläufe präsentiert werden.",
    )

    doc.add_heading("2. Projektüberblick", level=1)
    add_para(
        doc,
        "SYBAU ist eine digitale Fitness-Anwendung, die klassisches Workout-Tracking mit spielerischen Elementen verbindet. Ziel war es, Fortschritt nicht nur als Zahlenliste darzustellen, sondern als sichtbare Entwicklung über Level, Avatar, Belohnungen und soziale Vergleiche.",
    )
    add_para(
        doc,
        "Die Plattform besteht aus einer Website, einer Mobile-App und einem gemeinsamen Backend. Website und Mobile-App greifen auf dieselbe API und dieselbe Datenbank zu. Dadurch bleiben Konto, Fortschritt, Quests, Shop-Inhalte, Profil und Leaderboard auf beiden Plattformen konsistent.",
    )
    add_bullets(
        doc,
        [
            "Zielgruppe: Nutzer, die Training strukturierter und motivierender verfolgen wollen.",
            "Kernnutzen: schneller Trainingseintrag, sichtbarer Fortschritt und Belohnungssystem.",
            "Motivationsansatz: XP, Coins, Quests, Achievements, Avatar-Entwicklung und Leaderboard.",
            "Ergebnis: präsentierbarer Prototyp mit Web-App, Mobile-App und produktionsnaher Infrastruktur.",
        ],
    )

    doc.add_heading("3. Technische Umsetzung", level=1)
    add_para(
        doc,
        "Die technische Umsetzung wurde in drei Hauptbereiche aufgeteilt: Web-Frontend, Mobile-App und Backend. Das Backend stellt die zentrale Schnittstelle bereit und verwaltet Authentifizierung, Nutzerfortschritt, Workouts, Quests, Shop, Chests, Achievements, Freunde und Leaderboard.",
    )
    add_feature_table(
        doc,
        [
            ("Web-Frontend", "Vue 3, TypeScript und Vite. Die Website wird über Vercel bereitgestellt und ist unter https://sybau-fitness.vercel.app erreichbar."),
            ("Mobile-App", "Flutter und Dart. Die App nutzt dieselbe API wie die Website und bildet die wichtigsten Funktionen für mobile Nutzung ab."),
            ("Backend", "ASP.NET Core / C# mit Entity Framework Core. Das Backend läuft auf Render unter https://sybau-xll5.onrender.com."),
            ("Datenbank", "Neon Free PostgreSQL in Produktion. Lokal kann für Entwicklung SQLite verwendet werden."),
            ("Sicherheit", "JWT-basierte Authentifizierung, Admin-Schutz, Eingabevalidierung bei Login/Registrierung und Datenschutz-/Impressum-Seiten."),
            ("Deployment", "Frontend über Vercel, Backend über Render, Datenbank über Neon Free PostgreSQL."),
        ],
    )

    doc.add_heading("4. Umgesetzte Features", level=1)
    doc.add_heading("4.1 Website", level=2)
    add_feature_table(
        doc,
        [
            ("Landingpage", "Öffentliche Startseite mit Projektvorstellung, Call-to-Action zur Registrierung und responsiver Darstellung."),
            ("Auth", "Registrierung, Login, E-Mail-Prüfung, Passwort-Hinweise, Datenschutz-Checkbox und Schutz gegen Temp-Mail-Domains."),
            ("Dashboard", "Zentrale Übersicht mit Avatar, Level, XP, Coins, Streaks und direktem Fortschrittsfeedback."),
            ("Workouts", "50 Übungen mit Kategorien, Suchfeld, Schwierigkeit, Eingabeformat Zeit oder Einheiten sowie balancierten XP-/Coin-Belohnungen."),
            ("Quests", "Tägliche, wöchentliche und monatliche Quests mit Fortschritt, Belohnungen und live herunterzählendem Daily-Timer."),
            ("Shop", "Items, Booster, Chests, Coin-Pakete und Chest-Opening. Echtgeld-Angebote sind vorbereitet, aber als 'noch in Arbeit' gekennzeichnet."),
            ("Avatar", "Avatar-System mit sichtbarer Entwicklung und Item-/Booster-Bezug."),
            ("Profil", "Profilbild, Benutzername, privates Profil, Achievements, Statistiken und Weekly Activity Kalender."),
            ("Freunde", "Nutzersuche, Freundschaftsanfragen, öffentliche Profile und soziale Vergleiche."),
            ("Leaderboard", "Globale Rangliste mit Level- und XP-Vergleich."),
            ("Admin/Rechtliches", "Admin-Bereich zur Verwaltung von Inhalten sowie Impressum und Datenschutzerklärung."),
        ],
    )

    doc.add_heading("4.2 Mobile-App", level=2)
    add_para(
        doc,
        "Die Mobile-App wurde als Zusatzumfang umgesetzt. Sie war ursprünglich nicht als Hauptbestandteil des Projekts geplant, wurde aber fertiggestellt, um SYBAU auch unterwegs sinnvoll nutzbar zu machen.",
    )
    add_feature_table(
        doc,
        [
            ("Login und Konto", "Nutzer können sich mit demselben Konto wie auf der Website anmelden."),
            ("Dashboard", "Mobile Übersicht mit Avatar, Level, XP, Coins und Fortschritt."),
            ("Workouts", "Mobile Workout-Erfassung mit Suche, Kategorien und schneller Bedienung während oder nach dem Training."),
            ("Quests", "Tägliche, wöchentliche und monatliche Quests inklusive live herunterzählendem Daily-Timer."),
            ("Profil", "Profilansicht mit Achievements, Aktivitätskalender, Profilbild und Sichtbarkeitseinstellungen."),
            ("Freunde und Leaderboard", "Soziale Funktionen wie Freunde, öffentliche Profile und Rangliste sind mobil nutzbar."),
            ("Health-Sync", "Mobile Aktivitätsdaten können für Fortschritt und Quest-Fortschritt genutzt werden bzw. sind dafür vorbereitet."),
        ],
    )

    doc.add_heading("5. Soll-Ist-Vergleich", level=1)
    add_status_table(
        doc,
        [
            ("Account und Auth", "Erfüllt", "Registrierung, Login, Validierung, Datenschutz-Checkbox und geschützte Bereiche wurden umgesetzt."),
            ("Workout-System", "Erfüllt", "Die geplante Trainingserfassung wurde mit 50 Übungen, Suche, Kategorien und balancierten Belohnungen erweitert."),
            ("XP, Coins und Level", "Erfüllt", "Fortschritt wird berechnet, gespeichert und direkt im Header sichtbar gemacht."),
            ("Avatar-System", "Erfüllt", "Der Avatar ist als sichtbares Fortschrittselement integriert und mit dem Level-/Item-System verbunden."),
            ("Quests", "Erfüllt", "Daily, Weekly und Monthly Quests sind umgesetzt; Daily besitzt einen live laufenden Timer."),
            ("Shop und Chests", "Erfüllt", "Shop, Items, Booster und Chest-Opening sind umgesetzt. Echtgeldkäufe sind vorbereitet, aber ohne Zahlungssystem."),
            ("Freunde und Leaderboard", "Erfüllt", "Freundesfunktionen, öffentliche Profile und Ranglisten wurden umgesetzt."),
            ("Deployment", "Erfüllt", "Frontend, Backend und Datenbank sind produktionsnah über Vercel, Render und Neon Free erreichbar."),
            ("Mobile-App", "Zusatzumfang", "Eine vollständige Flutter-App wurde zusätzlich fertiggestellt und ist damit eine Scope-Erweiterung gegenüber dem ursprünglichen Web-Fokus."),
            ("Google-Fit/Web-Anbindung", "Angepasst", "Eine direkte Google-Fit-Anbindung in der Web-App wurde nicht als Kernfunktion umgesetzt. Stattdessen wurde der mobile Health-Sync-Ansatz verfolgt."),
            ("Echtgeld-Zahlung", "Vorbereitet", "Echtgeld-Angebote können dargestellt werden. Die echte Zahlungsabwicklung wurde bewusst nicht implementiert und zeigt aktuell 'noch in Arbeit'."),
        ],
    )
    add_para(
        doc,
        "Insgesamt wurde der geplante Kernumfang erreicht. Einzelne Punkte wurden im Projektverlauf angepasst, weil sich technische oder zeitliche Rahmenbedingungen verändert haben. Gleichzeitig wurde mit der Mobile-App ein zusätzlicher Liefergegenstand fertiggestellt, der den ursprünglichen Nutzen des Projekts erweitert.",
    )

    doc.add_heading("6. Finaler Statusreport", level=1)
    add_status_table(
        doc,
        [
            ("Web-Anwendung", "Fertig", "Die Website ist funktionsfähig, responsiv überarbeitet und für die Live-Demonstration geeignet."),
            ("Mobile-App", "Fertig", "Die Flutter-App deckt die wichtigsten Nutzerabläufe ab und nutzt dieselbe Backend-API."),
            ("Backend", "Fertig", "Die API stellt die zentralen Funktionen bereit und baut erfolgreich."),
            ("Datenbank", "Fertig", "Die Produktionsdaten werden in Neon Free PostgreSQL gespeichert."),
            ("Rechtliches", "Fertig", "Datenschutz und Impressum sind vorhanden und an die tatsächlich genutzten Dienste angepasst."),
            ("Qualität", "Präsentationsbereit", "Web-Build, Backend-Build, Flutter-Tests, Flutter-Analyse und iOS-Simulator-Build wurden geprüft."),
            ("Bekannte Einschränkung", "Akzeptiert", "Echte Zahlungsabwicklung für Echtgeld-Angebote ist nicht Teil der finalen Implementierung."),
        ],
    )
    add_callout(
        doc,
        "Finale Bewertung",
        "SYBAU ist als Projektabschluss präsentationsbereit. Die Kernfunktionen sind umgesetzt, der Zusatzumfang Mobile-App ist fertiggestellt und die technische Infrastruktur ist für eine Live-Produktpräsentation ausreichend stabil.",
    )

    doc.add_heading("7. Lessons Learned / Retrospective", level=1)
    doc.add_heading("Was gut funktioniert hat", level=2)
    add_bullets(
        doc,
        [
            "Die Aufteilung in Frontend, Backend und Mobile-App hat geholfen, parallel an mehreren Bereichen zu arbeiten.",
            "Gamification-Elemente wie XP, Coins, Quests und Chests haben das Projekt klarer und motivierender gemacht.",
            "Regelmäßiges Feedback führte zu vielen konkreten Verbesserungen bei Responsiveness, Datenschutz, Navigation und Nutzerführung.",
            "Die gemeinsame API für Website und Mobile-App war sinnvoll, weil beide Plattformen dieselben Daten verwenden können.",
        ],
    )
    doc.add_heading("Was herausfordernd war", level=2)
    add_bullets(
        doc,
        [
            "Die Abstimmung zwischen Web, Mobile, Backend und Datenbank wurde mit wachsendem Funktionsumfang komplexer.",
            "Deployment und Hosting mussten mehrfach angepasst werden, bis Frontend, Backend und Datenbank gut zusammenarbeiteten.",
            "Responsive Darstellung war besonders wichtig, da Website und Mobile-App beide sauber nutzbar sein sollten.",
            "Ein echtes Zahlungssystem wäre für den Projektumfang zu groß gewesen und wurde daher bewusst nur vorbereitet.",
        ],
    )
    doc.add_heading("Erkenntnisse für zukünftige Projekte", level=2)
    add_bullets(
        doc,
        [
            "Technische Infrastruktur sollte früh verbindlich festgelegt werden, damit spätere Hosting-Wechsel weniger Aufwand verursachen.",
            "Mobile Anforderungen sollten früh in der Planung berücksichtigt werden, auch wenn sie zuerst nur als Zusatzidee wirken.",
            "Rechtliche Texte und Datenschutz sollten nicht erst am Ende entstehen, sondern parallel zur tatsächlichen Datennutzung gepflegt werden.",
            "Kleine UI-Details wie Ladezustand, Tooltips, Timer und responsives Verhalten haben großen Einfluss auf die Präsentationsqualität.",
        ],
    )

    doc.add_heading("Abschluss", level=1)
    add_para(
        doc,
        "Das Projektziel wurde erreicht. SYBAU liegt als funktionsfähiges, präsentierbares System mit Website, Backend, Datenbank und Mobile-App vor. Besonders positiv ist, dass der Projektumfang über den ursprünglichen Web-Fokus hinaus erweitert wurde und mit der Mobile-App ein zusätzlicher praktischer Zugang zum System entstanden ist.",
    )


def remove_unused_tables_from_deleted_content(doc: Document) -> None:
    # python-docx keeps table wrappers in lists, but after XML body removal the body is authoritative.
    return None


def main() -> None:
    doc = Document(TEMPLATE)
    update_cover(doc)
    remove_old_report_content(doc)
    add_report_content(doc)
    update_footer_text(doc)
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)


if __name__ == "__main__":
    main()
